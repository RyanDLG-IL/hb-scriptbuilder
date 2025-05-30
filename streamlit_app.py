import streamlit as st
import os
import fitz as pymupdf
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import zipfile
import threading

# Import the prompt functions from the prompts directory
from prompts.activities_prompt import get_prompt as get_activities_prompt
from prompts.warmup_prompt import get_prompt as get_warmup_prompt
from prompts.summary_prompt import get_prompt as get_summary_prompt
from prompts.fact_check_prompt import get_prompt as get_fact_check_prompt
from prompts.dei_check_prompt import get_prompt as get_dei_check_prompt
from prompts.assessment_items_prompt import get_prompt as get_assessment_prompt
from prompts.script_visuals_tasks_prompt import get_prompt as get_visuals_tasks_prompt

# Load environment variables
load_dotenv()
genai_api_key = os.getenv("GEMINI_API_KEY")

if not genai_api_key:
    st.error("GEMINI_API_KEY not found. Please check your .env file.")
    st.stop()

genai.configure(api_key=genai_api_key)
# Model instances
activities_model = genai.GenerativeModel("gemini-1.5-flash")
warmup_model = genai.GenerativeModel("gemini-1.5-flash")
summary_model = genai.GenerativeModel("gemini-1.5-flash")
assessment_model = genai.GenerativeModel("gemini-1.5-flash")
fact_check_model = genai.GenerativeModel("gemini-1.5-flash")
dei_check_model = genai.GenerativeModel("gemini-1.5-flash")
assessment_fact_check_model = genai.GenerativeModel("gemini-1.5-flash")
assessment_dei_check_model = genai.GenerativeModel("gemini-1.5-flash")
visuals_tasks_model = genai.GenerativeModel("gemini-1.5-flash")

# Initialize session state variables if they don't exist
if 'lesson_blueprint' not in st.session_state:
    st.session_state.lesson_blueprint = None
if 'activities_output' not in st.session_state:
    st.session_state.activities_output = None
if 'warmup_output' not in st.session_state:
    st.session_state.warmup_output = None
if 'summary_output' not in st.session_state:
    st.session_state.summary_output = None
if 'assessment_items' not in st.session_state:
    st.session_state.assessment_items = None
if 'fact_check_output' not in st.session_state:
    st.session_state.fact_check_output = None
if 'dei_check_output' not in st.session_state:
    st.session_state.dei_check_output = None
if 'assessment_fact_check' not in st.session_state:
    st.session_state.assessment_fact_check = None
if 'assessment_dei_check' not in st.session_state:
    st.session_state.assessment_dei_check = None
if 'visuals_tasks_output' not in st.session_state:
    st.session_state.visuals_tasks_output = None
if 'has_generated' not in st.session_state:
    st.session_state.has_generated = False

# Add reset function
def reset_outputs():
    st.session_state.lesson_blueprint = None
    st.session_state.activities_output = None
    st.session_state.warmup_output = None
    st.session_state.summary_output = None
    st.session_state.assessment_items = None
    st.session_state.fact_check_output = None
    st.session_state.dei_check_output = None
    st.session_state.assessment_fact_check = None
    st.session_state.assessment_dei_check = None
    st.session_state.visuals_tasks_output = None
    st.session_state.has_generated = False

# Helper functions
def extract_text_from_pdf(file_content):
    try:
        with BytesIO(file_content) as pdf_buffer:
            doc = pymupdf.open(stream=pdf_buffer)
            text = ''
            for page in doc:
                text += page.get_text()
            return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file_content):
    try:
        with BytesIO(file_content) as docx_buffer:
            doc = Document(docx_buffer)
            return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    
    # Get file content
    file_content = uploaded_file.read()
    
    # Process based on file extension
    if uploaded_file.name.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_content)
    elif uploaded_file.name.lower().endswith(".docx"):
        return extract_text_from_docx(file_content)
    elif uploaded_file.name.lower().endswith((".txt", ".md")):
        return file_content.decode("utf-8")
    else:
        return f"Unsupported file format: {uploaded_file.name}"

def extract_text_from_path(file_path):
    try:
        if file_path.lower().endswith('.pdf'):
            doc = pymupdf.open(file_path)
            text = ''
            for page in doc:
                text += page.get_text()
            return text
        elif file_path.lower().endswith(('.txt', '.md')):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            return f"Unsupported file format: {file_path}"
    except Exception as e:
        return f"Error reading file ({file_path}): {str(e)}"

def load_reference_materials(folder_path):
    reference_texts = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        text = extract_text_from_path(file_path)
        reference_texts.append(f"Document: {filename}\n{text}\n\n")
    return "\n".join(reference_texts)

# Updated function to load specific reference materials (PDF and DOCX)
def load_dei_reference_materials(folder_path):
    dei_reference_texts = []
    target_files = ["Subject Specific Guidelines-Social Studies", "DEI Content Authoring Guidelines"]
    
    for filename in os.listdir(folder_path):
        # Check if the filename contains any of the target file names
        if any(target_file in filename for target_file in target_files):
            file_path = os.path.join(folder_path, filename)
            text = extract_text_from_path(file_path)
            dei_reference_texts.append(f"Document: {filename}\n{text}\n\n")
    
    return "\n".join(dei_reference_texts)

# Load reference materials
reference_materials_folder = 'reference_materials'
reference_content = load_reference_materials(reference_materials_folder)

# Load DEI-specific reference materials
dei_reference_content = load_dei_reference_materials(reference_materials_folder)

# Streamlit page setup
st.set_page_config(page_title="Lesson Script Generator", layout="wide")
st.title("Lesson Script Generator")

# Input field for the lesson blueprint only
st.markdown("Input your lesson blueprint below. Assessment items will be automatically generated.")

# Lesson Blueprint input
st.subheader("Lesson Blueprint")

blueprint_input_method = st.radio(
    "How would you like to provide the Lesson Blueprint?",
    options=["Paste text", "Upload file"],
    horizontal=True,
    key="blueprint_method"
)

lesson_blueprint_text = ""
if blueprint_input_method == "Paste text":
    lesson_blueprint_text = st.text_area("Paste your Lesson Blueprint here:", height=300, 
                                help="Paste the complete lesson blueprint created with the blueprint tool")
else:
    uploaded_blueprint = st.file_uploader("Upload your Lesson Blueprint file", 
                                        type=["txt", "pdf", "docx", "md"],
                                        help="Upload the lesson blueprint file (PDF, Word, Text)")
    if uploaded_blueprint is not None:
        lesson_blueprint_text = extract_text_from_file(uploaded_blueprint)
        st.success(f"Successfully loaded blueprint from {uploaded_blueprint.name}")
        # Preview the extracted text
        with st.expander("Preview extracted blueprint text"):
            st.text(lesson_blueprint_text[:500] + ("..." if len(lesson_blueprint_text) > 500 else ""))

# New functions for generating scripts
def create_activities_script(blueprint):
    # Use the activities prompt to generate the activities script
    prompt = get_activities_prompt(blueprint)
    
    # Add DEI considerations
    prompt += f"""
    
    ### Additional DEI Considerations
    Please ensure your activities script follows these DEI guidelines:
    {dei_reference_content}
    """
    
    try:
        response = activities_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating activities script: {str(e)}"

def create_warmup_script(blueprint, activities_script):
    # Use the warmup prompt to generate the warmup script
    # Extract metadata from blueprint for the prompt
    prompt = f"""
    # Warmup Script Generation
    
    ## Lesson Blueprint:
    {blueprint}
    
    ## Activities Script:
    {activities_script}
    
    {get_warmup_prompt('', '', '', '', activities_script)}
    
    ### Additional DEI Considerations
    Please ensure your warmup script follows these DEI guidelines:
    {dei_reference_content}
    
    Note: Extract all necessary metadata (unit title, lesson title, learning objectives, standard codes) 
    directly from the lesson blueprint provided. Use that extracted information to create an appropriate warmup script.
    """
    
    try:
        response = warmup_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating warmup script: {str(e)}"

def create_summary_script(blueprint, warmup_script, activities_script):
    # Use the summary prompt to generate the summary script
    # Extract metadata from blueprint for the prompt
    prompt = f"""
    # Summary Script Generation
    
    ## Lesson Blueprint:
    {blueprint}
    
    ## Warmup Script:
    {warmup_script}
    
    ## Activities Script:
    {activities_script}
    
    {get_summary_prompt('', '', '', '', '', warmup_script, activities_script)}
    
    ### Additional DEI Considerations
    Please ensure your summary script follows these DEI guidelines:
    {dei_reference_content}
    
    Note: Extract all necessary metadata (unit title, lesson title, lesson question, learning objectives, standard codes) 
    directly from the lesson blueprint provided. Use that extracted information to create an appropriate summary script.
    """
    
    try:
        response = summary_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating summary script: {str(e)}"

# New function to generate assessment items
def create_assessment_items(blueprint, activities_script, warmup_script, summary_script, fact_check, dei_check):
    # Use the assessment prompt to generate assessment items
    prompt = get_assessment_prompt(
        blueprint, 
        activities_script, 
        warmup_script, 
        summary_script, 
        fact_check, 
        dei_check
    )
    
    try:
        response = assessment_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating assessment items: {str(e)}"

# Functions for fact checking and DEI checking
def create_fact_check(blueprint, activities_script, warmup_script, summary_script):
    # Call the fact check prompt function with all parameters
    prompt = get_fact_check_prompt(blueprint, "", activities_script, warmup_script, summary_script)
    
    try:
        response = fact_check_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating fact check: {str(e)}"

def create_dei_check(blueprint, activities_script, warmup_script, summary_script):
    # Updated to pass parameters to the get_dei_check_prompt function
    prompt = get_dei_check_prompt(
        blueprint, 
        "", 
        activities_script, 
        warmup_script, 
        summary_script, 
        dei_reference_content
    )
    
    try:
        response = dei_check_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating DEI check: {str(e)}"

# Functions for assessment fact checking and DEI checking
def create_assessment_fact_check(blueprint, assessment_items):
    # Create a fact check for assessment items
    prompt = f"""
    ## Context ##
    You are an expert educational content reviewer specializing in fact-checking assessment items for social studies. 
    Your task is to carefully examine the provided assessment items for factual accuracy.

    ## Objective ##
    Review the provided assessment items to identify and report on any potential factual inaccuracies, 
    misleading information, or content that requires verification.

    ## Response ##
    Provide a structured fact-check report with the following sections:

    1. **Overall Accuracy Assessment**
       - Overall accuracy rating (choose one): 
         ✅ Highly Accurate (No factual errors found)
         ⚠️ Generally Accurate (Contains minor factual issues)
         ❌ Needs Significant Revision (Contains major factual errors)
       - Brief summary of the assessment items' factual reliability

    2. **Specific Factual Issues** (if any)
       - For each potential factual issue, provide:
         
         **Original Question/Item**: [original text]
         
         - Fact-checking Report: [description of the issue with citations to reliable sources]
         
         **Revised Question/Item**: [revised version]
         
    ### Content to Review:
    
    #### Lesson Blueprint:
    {blueprint}
    
    #### Assessment Items:
    {assessment_items}
    """
    
    try:
        response = assessment_fact_check_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating assessment fact check: {str(e)}"

def create_assessment_dei_check(blueprint, assessment_items):
    # Create a DEI check for assessment items
    prompt = f"""
    ## Context ##
    You are a DEI content review assistant trained to support the development of inclusive, accurate, 
    and respectful Social Studies assessment items. Your review must follow Imagine Learning's DEI Content 
    Development Guidelines, which emphasize diverse representation, critical analysis, factual grounding, 
    and awareness of historical complexity.

    ## Objective ##
    Review the provided assessment items for Social Studies and determine how well they align with 
    DEI expectations for the Social Studies space.

    ## Response ##
    Return your review using this structure:

    **Overall Alignment Rating**:  
    [Choose one: ✅ Strong Alignment | ⚠️ Moderate Alignment | ❌ Needs Significant Improvement]

    **Key Findings**:  
    - [List key strengths or concerns in the assessment items]

    **Specific Recommendations**
    
    For each issue identified, provide:
    
    **Original Question/Item**: [original text]
    
    **Issue**: [description of the DEI concern]
    
    **Recommended Revision**: [suggested revision]
    
    **Rationale**: [brief explanation based on DEI principles]

    ## Reference Materials ##
    {dei_reference_content}
    
    ### Content to Review:
    
    #### Lesson Blueprint:
    {blueprint}
    
    #### Assessment Items:
    {assessment_items}
    """
    
    try:
        response = assessment_dei_check_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating assessment DEI check: {str(e)}"

# Add function to create visuals and tasks suggestions
def create_visuals_tasks_suggestions(blueprint, activities_script, warmup_script, summary_script, 
                                     fact_check, dei_check, assessment_items):
    prompt = get_visuals_tasks_prompt(
        blueprint, 
        activities_script, 
        warmup_script, 
        summary_script,
        fact_check,
        dei_check,
        assessment_items
    )
    
    try:
        response = visuals_tasks_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating visuals and tasks suggestions: {str(e)}"

# Function to create a Word document
def create_word_doc(title, content):
    doc = Document()
    # Add title with formatting
    title_heading = doc.add_heading(title, level=1)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Split content into paragraphs
    paragraphs = content.split('\n')
    
    # Flag to track if we're inside a table
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if not para.strip():
            i += 1
            continue
        
        # Check for table start (line contains | and the next line has dashes and |)
        if '|' in para and i + 1 < len(paragraphs) and set(paragraphs[i+1].replace('|', '')).issubset({'-', ' '}):
            in_table = True
            table_rows = []
            
            # Add header row
            header_cells = [cell.strip() for cell in para.split('|')[1:-1]] if para.strip().startswith('|') else [cell.strip() for cell in para.split('|')]
            table_rows.append(header_cells)
            
            # Skip the separator row
            i += 2  # Move past header and separator lines
            
            # Process table rows
            while i < len(paragraphs) and '|' in paragraphs[i]:
                row = paragraphs[i]
                row_cells = [cell.strip() for cell in row.split('|')[1:-1]] if row.strip().startswith('|') else [cell.strip() for cell in row.split('|')]
                table_rows.append(row_cells)
                i += 1
            
            # Create the Word table
            if table_rows:
                max_cols = max(len(row) for row in table_rows)
                word_table = doc.add_table(rows=len(table_rows), cols=max_cols)
                word_table.style = 'Table Grid'
                
                # Fill the table
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < max_cols:  # Ensure we don't exceed column count
                            cell = word_table.cell(row_idx, col_idx)
                            cell.text = cell_text
                            
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                # Add space after table
                doc.add_paragraph()
            
            in_table = False
            continue
            
        # Check if paragraph is a header (markdown style)
        header_match = re.match(r'^#+\s+(.+)$', para)
        if header_match:
            header_text = header_match.group(1)
            level = min(len(re.match(r'^#+', para).group(0)), 6)
            doc.add_heading(header_text, level=level)
            i += 1
            continue
            
        # Check if paragraph is a list item
        if para.strip().startswith('- ') or para.strip().startswith('* '):
            doc.add_paragraph(para.strip()[2:], style='List Bullet')
            i += 1
            continue
            
        # Handle bolded text with markdown ** or __
        bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
        if re.search(bold_pattern, para):
            p = doc.add_paragraph()
            parts = re.split(bold_pattern, para)
            for j, part in enumerate(parts):
                if part:  # Skip empty parts
                    if j % 3 == 1 or j % 3 == 2:  # Every third part matches the bold pattern
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            i += 1
            continue
            
        # Regular paragraph
        doc.add_paragraph(para)
        i += 1
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Add a function to create a combined lesson script document
def create_combined_script_doc(warmup_script, activities_script, summary_script):
    doc = Document()
    # Add title with formatting
    title_heading = doc.add_heading("Complete Lesson Script", level=1)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add each section with proper formatting
    
    # Warmup Section
    doc.add_heading("Warmup Script", level=2)
    warmup_paragraphs = warmup_script.split('\n')
    for para in warmup_paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    # Add a page break between sections
    doc.add_page_break()
    
    # Activities Section
    doc.add_heading("Activities Script", level=2)
    activities_paragraphs = activities_script.split('\n')
    for para in activities_paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    # Add a page break between sections
    doc.add_page_break()
    
    # Summary Section
    doc.add_heading("Summary Script", level=2)
    summary_paragraphs = summary_script.split('\n')
    for para in summary_paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Update the zip creation function to include visuals/tasks
def create_zip_with_all_docs(combined_script_doc, script_reports_doc, assessment_doc, 
                            assessment_reports_doc, visuals_tasks_doc):
    # Create a ZIP file in memory
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
        # Add all generated content to the zip file
        zip_file.writestr("Lesson_Script.docx", combined_script_doc.getvalue())
        zip_file.writestr("Script_DEI_Fact_Check_Reports.docx", script_reports_doc.getvalue())
        zip_file.writestr("Assessment_Items.docx", assessment_doc.getvalue())
        zip_file.writestr("Assessment_DEI_Fact_Check_Reports.docx", assessment_reports_doc.getvalue())
        zip_file.writestr("Visual_and_Task_Suggestions.docx", visuals_tasks_doc.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

# Add sidebar buttons for generation and reset
st.sidebar.title("Controls")
generate_button = st.sidebar.button("Generate Scripts", use_container_width=True, key="generate_btn")
reset_button = st.sidebar.button("Reset Tool", on_click=reset_outputs, use_container_width=True, key="reset_btn")

# Add a status section in sidebar
st.sidebar.markdown("---")
status_container = st.sidebar.container()

# Add a separator between status and download options (initially hidden)
download_separator = st.sidebar.empty()

# Update content generation workflow
if generate_button:
    if not lesson_blueprint_text.strip():
        st.warning("Please provide a lesson blueprint to generate scripts.")
    else:
        # Store the inputs
        st.session_state.lesson_blueprint = lesson_blueprint_text
        
        # Step 1: Generate activities script
        with status_container.status("Creating activities script"):
            st.session_state.activities_output = create_activities_script(lesson_blueprint_text)
        
        # Step 2: Generate warmup script
        with status_container.status("Creating warmup script"):
            st.session_state.warmup_output = create_warmup_script(
                lesson_blueprint_text,
                st.session_state.activities_output
            )
        
        # Step 3: Generate summary script
        with status_container.status("Creating summary script"):
            st.session_state.summary_output = create_summary_script(
                lesson_blueprint_text,
                st.session_state.warmup_output,
                st.session_state.activities_output
            )
        
        # Step 4: Fact check and DEI check the scripts
        with status_container.status("Reviewing scripts for accuracy and inclusion"):
            # Store local variables for the threads to work with
            blueprint_content = st.session_state.lesson_blueprint
            activities_content = st.session_state.activities_output
            warmup_content = st.session_state.warmup_output
            summary_content = st.session_state.summary_output
            
            # Use lists to store results (mutable objects that can be modified by threads)
            fact_check_result = [None]
            dei_check_result = [None]

            # Define thread functions that don't access session state
            def run_fact_check():
                fact_check_result[0] = create_fact_check(
                    blueprint_content,
                    activities_content, 
                    warmup_content,
                    summary_content
                )

            def run_dei_check():
                dei_check_result[0] = create_dei_check(
                    blueprint_content,
                    activities_content, 
                    warmup_content,
                    summary_content
                )

            # Create and start both threads
            fact_check_thread = threading.Thread(target=run_fact_check)
            dei_check_thread = threading.Thread(target=run_dei_check)
            fact_check_thread.start()
            dei_check_thread.start()

            # Wait for both to complete
            fact_check_thread.join()
            dei_check_thread.join()

            # Now that threads are done, save results to session state in main thread
            st.session_state.fact_check_output = fact_check_result[0]
            st.session_state.dei_check_output = dei_check_result[0]
        
        # Step 5: Generate assessment items based on all previous content
        with status_container.status("Creating assessment items"):
            st.session_state.assessment_items = create_assessment_items(
                st.session_state.lesson_blueprint,
                st.session_state.activities_output,
                st.session_state.warmup_output,
                st.session_state.summary_output,
                st.session_state.fact_check_output,
                st.session_state.dei_check_output
            )
        
        # Step 6: Fact check and DEI check the assessment items
        with status_container.status("Reviewing assessment items"):
            # Store variables for assessment check
            blueprint_content = st.session_state.lesson_blueprint
            assessment_content = st.session_state.assessment_items
            
            # Use lists to store assessment check results
            assessment_fact_check_result = [None]
            assessment_dei_check_result = [None]
            
            def run_assessment_fact_check():
                assessment_fact_check_result[0] = create_assessment_fact_check(
                    blueprint_content,
                    assessment_content
                )
                
            def run_assessment_dei_check():
                assessment_dei_check_result[0] = create_assessment_dei_check(
                    blueprint_content,
                    assessment_content
                )
                
            # Create and start assessment check threads
            assessment_fact_thread = threading.Thread(target=run_assessment_fact_check)
            assessment_dei_thread = threading.Thread(target=run_assessment_dei_check)
            assessment_fact_thread.start()
            assessment_dei_thread.start()
            
            # Wait for both to complete
            assessment_fact_thread.join()
            assessment_dei_thread.join()
            
            # Save results to session state
            st.session_state.assessment_fact_check = assessment_fact_check_result[0]
            st.session_state.assessment_dei_check = assessment_dei_check_result[0]
        
        # Step 7: Generate visuals and tasks suggestions
        with status_container.status("Creating visual and task suggestions"):
            st.session_state.visuals_tasks_output = create_visuals_tasks_suggestions(
                st.session_state.lesson_blueprint,
                st.session_state.activities_output,
                st.session_state.warmup_output,
                st.session_state.summary_output,
                st.session_state.fact_check_output,
                st.session_state.dei_check_output,
                st.session_state.assessment_items
            )
        
        # All content is now generated - display final success message
        status_container.success("All content successfully generated!")
        
        # Set flag that content has been generated
        st.session_state.has_generated = True
        
        # Show separator for download section
        download_separator.markdown("---")
        
        # Use the current rerun method
        st.rerun()

# Display outputs if they exist
if st.session_state.has_generated:
    # Create combined lesson script document
    combined_script_doc = create_combined_script_doc(
        st.session_state.warmup_output,
        st.session_state.activities_output,
        st.session_state.summary_output
    )
    
    # Create script reports document
    script_reports_doc = create_word_doc(
        "Script DEI and Fact-check Reports", 
        f"## Fact Check Report\n\n{st.session_state.fact_check_output}\n\n## DEI Check Report\n\n{st.session_state.dei_check_output}"
    )
    
    # Create assessment documents
    assessment_doc = create_word_doc("Assessment Items", st.session_state.assessment_items)
    assessment_reports_doc = create_word_doc(
        "Assessment DEI and Fact-check Reports",
        f"## Assessment Fact Check Report\n\n{st.session_state.assessment_fact_check}\n\n## Assessment DEI Check Report\n\n{st.session_state.assessment_dei_check}"
    )
    
    # Add visuals and tasks document
    visuals_tasks_doc = create_word_doc(
        "Visual and Task Suggestions", 
        st.session_state.visuals_tasks_output
    )
    
    # Create zip file with all content including new document
    all_docs_zip = create_zip_with_all_docs(
        combined_script_doc, 
        script_reports_doc,
        assessment_doc,
        assessment_reports_doc,
        visuals_tasks_doc
    )
    
    # SIDEBAR DOWNLOAD OPTIONS
    st.sidebar.markdown("### Download Options")
    
    # Make the Download All button visually distinct
    st.sidebar.markdown("""
    <style>
    div.stDownloadButton > button {
        background-color: #4CAF50 !important;
        color: white !important;
        font-weight: bold !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Download All button (with unique key to apply style)
    st.sidebar.download_button(
        label="💾 Download All Content",
        data=all_docs_zip,
        file_name="Lesson_Materials.zip",
        mime="application/zip",
        key="all_docs_download",
        use_container_width=True
    )
    
    # Reset the style for other buttons
    st.sidebar.markdown("""
    <style>
    div.stDownloadButton > button {
        background-color: inherit !important;
        color: inherit !important;
        font-weight: normal !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Individual download options
    st.sidebar.download_button(
        label="Lesson Script",
        data=combined_script_doc,
        file_name="Lesson_Script.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="combined_script_download",
        use_container_width=True
    )
    
    st.sidebar.download_button(
        label="Assessment Items",
        data=assessment_doc,
        file_name="Assessment_Items.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="assessment_download",
        use_container_width=True
    )
    
    st.sidebar.download_button(
        label="Script DEI & Fact Check",
        data=script_reports_doc,
        file_name="Script_DEI_Fact_Check_Reports.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="script_reports_download",
        use_container_width=True
    )
    
    st.sidebar.download_button(
        label="Assessment DEI & Fact Check",
        data=assessment_reports_doc,
        file_name="Assessment_DEI_Fact_Check_Reports.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="assessment_reports_download",
        use_container_width=True
    )
    
    st.sidebar.download_button(
        label="Visual & Task Suggestions",
        data=visuals_tasks_doc,
        file_name="Visual_and_Task_Suggestions.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="visuals_tasks_download",
        use_container_width=True
    )
    
    st.markdown("---")
    st.header("Generated Content")
    
    # CONTENT IN TABS
    activities_tab, warmup_tab, summary_tab, assessment_tab, script_fact_tab, script_dei_tab, assessment_fact_tab, assessment_dei_tab, visuals_tasks_tab = st.tabs([
        "Activities Script", 
        "Warmup Script", 
        "Summary Script",
        "Assessment Items",
        "Script Fact Check", 
        "Script DEI Check",
        "Assessment Fact Check",
        "Assessment DEI Check",
        "Visual & Task Suggestions"
    ])
    
    # Display content in each tab
    with activities_tab:
        st.markdown("## Activities Script")
        st.write(st.session_state.activities_output)
    
    with warmup_tab:
        st.markdown("## Warmup Script")
        st.write(st.session_state.warmup_output)
    
    with summary_tab:
        st.markdown("## Summary Script")
        st.write(st.session_state.summary_output)
        
    with assessment_tab:
        st.markdown("## Assessment Items")
        st.write(st.session_state.assessment_items)
    
    with script_fact_tab:
        st.markdown("## Script Fact Check Report")
        st.write(st.session_state.fact_check_output)
    
    with script_dei_tab:
        st.markdown("## Script DEI Check Report")
        st.write(st.session_state.dei_check_output)
        
    with assessment_fact_tab:
        st.markdown("## Assessment Fact Check Report")
        st.write(st.session_state.assessment_fact_check)
        
    with assessment_dei_tab:
        st.markdown("## Assessment DEI Check Report")
        st.write(st.session_state.assessment_dei_check)
    
    with visuals_tasks_tab:
        st.markdown("## Visual and Task Suggestions")
        st.write(st.session_state.visuals_tasks_output)